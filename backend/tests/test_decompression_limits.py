"""압축 폭탄 방어.

업로드는 로그인 없이 가능하고(미리보기 용도) 머신은 1GB다. 업로드 크기만
제한하면 10MB짜리 DOCX/HWP 하나가 압축 해제 후 몇 GB가 되어 프로세스를
죽일 수 있었다.
"""
import zipfile
import zlib

import pytest
from fastapi import HTTPException

import text_processing
from text_processing import extract_text


def _zip_with(tmp_path, entries: dict[str, bytes], name="bomb.docx"):
    path = tmp_path / name
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        for entry_name, payload in entries.items():
            archive.writestr(entry_name, payload)
    return path


def test_docx_that_explodes_when_unpacked_is_rejected(tmp_path, monkeypatch):
    monkeypatch.setattr(text_processing, "MAX_DECOMPRESSED_BYTES", 1024 * 1024)
    # 0으로 채운 2MB는 압축하면 몇 KB다 — 업로드 상한은 가뿐히 통과한다.
    path = _zip_with(tmp_path, {"word/document.xml": b"\0" * (2 * 1024 * 1024)})
    assert path.stat().st_size < 100 * 1024, "압축된 파일 자체는 작아야 재현이 성립한다"

    with pytest.raises(HTTPException) as exc:
        extract_text(str(path), "bomb.docx")

    assert exc.value.status_code == 413


def test_docx_with_too_many_entries_is_rejected(tmp_path, monkeypatch):
    """크기가 아니라 항목 수로 터뜨리는 쪽도 막는다."""
    monkeypatch.setattr(text_processing, "MAX_ARCHIVE_ENTRIES", 5)
    path = _zip_with(tmp_path, {f"f{i}.xml": b"x" for i in range(10)})

    with pytest.raises(HTTPException) as exc:
        extract_text(str(path), "bomb.docx")

    assert exc.value.status_code == 413


def test_corrupt_docx_reports_bad_file_not_bomb(tmp_path):
    path = tmp_path / "broken.docx"
    path.write_bytes(b"this is not a zip at all")

    with pytest.raises(HTTPException) as exc:
        extract_text(str(path), "broken.docx")

    assert exc.value.status_code == 400


def test_normal_docx_still_parses(tmp_path):
    """방어가 정상 문서를 막으면 안 된다."""
    import docx

    document = docx.Document()
    document.add_paragraph("첫 문단입니다.")
    document.add_paragraph("둘째 문단입니다.")
    path = tmp_path / "normal.docx"
    document.save(str(path))

    assert "첫 문단입니다." in extract_text(str(path), "normal.docx")


def test_hwp_stream_decompression_is_capped():
    """zlib.decompress()는 출력 크기 제한이 없어, 작은 스트림이 몇 GB로
    부풀어도 그대로 메모리에 올린다."""
    compressor = zlib.compressobj(9, zlib.DEFLATED, -15)
    bomb = compressor.compress(b"\0" * (4 * 1024 * 1024)) + compressor.flush()
    assert len(bomb) < 50 * 1024

    with pytest.raises(HTTPException) as exc:
        text_processing._decompress_limited(bomb, remaining=1024 * 1024)

    assert exc.value.status_code == 413


class _FakeOleFile:
    """olefile은 읽기 전용이라 진짜 HWP를 만들 수 없다. extract_hwp_text가
    보는 최소 구조(FileHeader + BodyText)만 흉내낸다."""

    def __init__(self, body: bytes):
        self._body = body

    def listdir(self):
        return [["FileHeader"], ["BodyText", "Section0"]]

    def openstream(self, name):
        import io

        if name == "FileHeader":
            header = bytearray(64)
            header[36] = 1  # 압축됨 플래그
            return io.BytesIO(bytes(header))
        return io.BytesIO(self._body)


def test_hwp_extraction_never_uses_unbounded_decompress(monkeypatch):
    """헬퍼만 안전해도 extract_hwp_text가 그걸 안 쓰면 의미가 없다.

    "결과가 413이다"로는 이걸 잡을 수 없다 — zlib.decompress()로 되돌려도
    푼 뒤의 크기 검사에 걸려 똑같이 413이 나온다. 문제는 결과가 아니라
    **거기 닿기 전에 이미 몇 GB를 메모리에 올렸다는 것**이고, 그건 메모리로
    관측할 수 없으니 "무제한 API를 호출하지 않는다"로 고정한다.
    """
    monkeypatch.setattr(text_processing, "MAX_DECOMPRESSED_BYTES", 1024 * 1024)
    compressor = zlib.compressobj(9, zlib.DEFLATED, -15)
    bomb = compressor.compress(b"\0" * (4 * 1024 * 1024)) + compressor.flush()

    import olefile

    monkeypatch.setattr(olefile, "OleFileIO", lambda _path: _FakeOleFile(bomb))

    def unbounded_is_forbidden(*_args, **_kwargs):
        raise AssertionError("zlib.decompress()는 출력 크기 제한이 없어 쓰면 안 된다")

    monkeypatch.setattr(text_processing.zlib, "decompress", unbounded_is_forbidden)

    with pytest.raises(HTTPException) as exc:
        text_processing.extract_hwp_text("ignored.hwp")

    # 413이어야 한다 — 일반 예외 처리에 걸려 400 "해석 실패"로 둔갑하면
    # 왜 거부됐는지 사용자도 로그도 알 수 없다.
    assert exc.value.status_code == 413


def test_hwp_decompression_returns_data_under_the_cap():
    payload = "안녕하세요 " * 100
    compressor = zlib.compressobj(9, zlib.DEFLATED, -15)
    packed = compressor.compress(payload.encode("utf-16le")) + compressor.flush()

    unpacked = text_processing._decompress_limited(packed, remaining=1024 * 1024)

    assert unpacked.decode("utf-16le") == payload
